"""
SRS Generator – chuyển kết quả phân tích thành file DOCX chuẩn.
"""
import os
import re
import frappe
from frappe.utils import get_files_path, now_datetime
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── Color palette ─────────────────────────────────────────────────────────────
C_TITLE   = RGBColor(0,  63, 127)
C_H1      = RGBColor(0,  70, 127)
C_H2      = RGBColor(31, 73, 125)
C_TH_BG   = "003F7F"
C_NEW     = RGBColor(192,  0,  0)
C_PARTIAL = RGBColor(192, 96,  0)
C_DONE    = RGBColor(  0, 112, 54)
C_GRAY    = RGBColor(120, 120, 120)


def generate(project_name: str) -> str:
	"""
	Sinh file SRS DOCX cho project và trả về file_url trong Frappe.
	"""
	from mbwnext_reqpilot.mbwnext_reqpilot.services.claude_client import generate_srs_prompt

	project = frappe.get_doc("SRS Project", project_name)

	# 1. Lấy SRS text từ Claude
	srs_markdown = generate_srs_prompt(project_name)

	# 2. Build DOCX
	doc = _build_docx(project, srs_markdown)

	# 3. Lưu file
	file_name = f"SRS_{project_name}_{now_datetime().strftime('%Y%m%d_%H%M')}.docx"
	site_files = get_files_path()
	file_path = os.path.join(site_files, file_name)
	doc.save(file_path)

	# 4. Tạo File record trong Frappe
	file_doc = frappe.get_doc({
		"doctype": "File",
		"file_name": file_name,
		"file_url": f"/files/{file_name}",
		"is_private": 0,
		"attached_to_doctype": "SRS Project",
		"attached_to_name": project_name,
		"attached_to_field": "output_srs"
	})
	file_doc.insert(ignore_permissions=True)

	project.output_srs = file_doc.file_url
	project.status = "Completed"
	project.save(ignore_permissions=True)
	frappe.db.commit()

	return file_doc.file_url


def _build_docx(project, srs_markdown: str) -> Document:
	doc = Document()

	# Page margins
	section = doc.sections[0]
	section.top_margin = Cm(2)
	section.bottom_margin = Cm(2)
	section.left_margin = Cm(2.5)
	section.right_margin = Cm(2.5)

	_cover_page(doc, project)
	doc.add_page_break()
	_render_markdown(doc, srs_markdown)
	_requirements_appendix(doc, project)

	return doc


def _cover_page(doc, project):
	doc.add_paragraph()
	p = doc.add_paragraph()
	p.alignment = WD_ALIGN_PARAGRAPH.CENTER
	r = p.add_run("TÀI LIỆU ĐẶC TẢ YÊU CẦU PHẦN MỀM")
	_font(r, 20, True, color=C_TITLE)

	p = doc.add_paragraph()
	p.alignment = WD_ALIGN_PARAGRAPH.CENTER
	r = p.add_run("Software Requirements Specification (SRS)")
	_font(r, 13, italic=True, color=C_GRAY)

	doc.add_paragraph()

	p = doc.add_paragraph()
	p.alignment = WD_ALIGN_PARAGRAPH.CENTER
	r = p.add_run(f"Dự án: {project.project_name}")
	_font(r, 15, True)

	p = doc.add_paragraph()
	p.alignment = WD_ALIGN_PARAGRAPH.CENTER
	r = p.add_run(f"Khách hàng: {project.customer or 'N/A'}  |  App: {project.custom_app_name or 'N/A'}")
	_font(r, 12)

	doc.add_paragraph()
	purchased = ", ".join(row.app_name for row in project.base_apps if row.included)

	tbl = doc.add_table(rows=4, cols=2)
	tbl.style = "Table Grid"
	tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
	meta = [
		("Phiên bản", "1.0"),
		("Ngày lập", now_datetime().strftime("%d/%m/%Y")),
		("Base Apps", purchased or "N/A"),
		("BA phụ trách", project.created_by_user or frappe.session.user),
	]
	for i, (k, v) in enumerate(meta):
		tbl.rows[i].cells[0].text = ""
		r = tbl.rows[i].cells[0].paragraphs[0].add_run(k)
		_font(r, 11, True)
		tbl.rows[i].cells[1].text = ""
		r2 = tbl.rows[i].cells[1].paragraphs[0].add_run(v)
		_font(r2, 11)


def _render_markdown(doc, markdown_text: str):
	"""Chuyển Markdown thành paragraphs/headings trong DOCX."""
	lines = markdown_text.split("\n")
	in_table = False
	table_rows = []

	for line in lines:
		stripped = line.rstrip()

		# Headings
		if stripped.startswith("### "):
			_flush_table(doc, table_rows); table_rows = []; in_table = False
			p = doc.add_heading(level=3); p.clear()
			r = p.add_run(stripped[4:]); _font(r, 11, True, color=C_H2)
		elif stripped.startswith("## "):
			_flush_table(doc, table_rows); table_rows = []; in_table = False
			p = doc.add_heading(level=2); p.clear()
			r = p.add_run(stripped[3:]); _font(r, 12, True, color=C_H1)
		elif stripped.startswith("# "):
			_flush_table(doc, table_rows); table_rows = []; in_table = False
			p = doc.add_heading(level=1); p.clear()
			r = p.add_run(stripped[2:]); _font(r, 14, True, color=C_TITLE)
		# Table row
		elif stripped.startswith("|"):
			parts = [c.strip() for c in stripped.strip("|").split("|")]
			if all(set(c) <= {"-", " ", ":"} for c in parts):
				continue  # separator row
			table_rows.append(parts)
			in_table = True
		# Bullet
		elif stripped.startswith("- ") or stripped.startswith("* "):
			_flush_table(doc, table_rows); table_rows = []; in_table = False
			p = doc.add_paragraph(style="List Bullet")
			r = p.add_run(stripped[2:]); _font(r, 11)
		# Code block (skip markers)
		elif stripped.startswith("```"):
			_flush_table(doc, table_rows); table_rows = []; in_table = False
		# Normal paragraph
		elif stripped:
			_flush_table(doc, table_rows); table_rows = []; in_table = False
			p = doc.add_paragraph()
			r = p.add_run(stripped); _font(r, 11)
		else:
			if in_table:
				_flush_table(doc, table_rows); table_rows = []; in_table = False

	_flush_table(doc, table_rows)


def _flush_table(doc, rows: list):
	if not rows:
		return
	ncols = max(len(r) for r in rows)
	tbl = doc.add_table(rows=1, cols=ncols)
	tbl.style = "Table Grid"

	# Header row
	hdr = tbl.rows[0].cells
	for i, val in enumerate(rows[0]):
		if i >= ncols:
			break
		hdr[i].text = ""
		r = hdr[i].paragraphs[0].add_run(val)
		_font(r, 10, True, color=RGBColor(255, 255, 255))
		_shd(hdr[i], C_TH_BG)

	# Data rows
	for ri, row in enumerate(rows[1:]):
		cells = tbl.add_row().cells
		bg = "EBF3FB" if ri % 2 == 0 else "FFFFFF"
		for ci, val in enumerate(row):
			if ci >= ncols:
				break
			cells[ci].text = ""
			r = cells[ci].paragraphs[0].add_run(val)
			_font(r, 10)
			if bg != "FFFFFF":
				_shd(cells[ci], bg)


def _requirements_appendix(doc, project):
	doc.add_page_break()
	p = doc.add_heading(level=1); p.clear()
	r = p.add_run("Phụ lục – Bảng tổng hợp Requirements")
	_font(r, 14, True, color=C_TITLE)

	if not project.requirements:
		doc.add_paragraph("Chưa có requirements được phân tích.")
		return

	tbl = doc.add_table(rows=1, cols=6)
	tbl.style = "Table Grid"
	headers = ["Mã YC", "Yêu cầu", "Trạng thái", "App", "Doctype", "Effort (ngày)"]
	hdr = tbl.rows[0].cells
	for i, h in enumerate(headers):
		hdr[i].text = ""
		r = hdr[i].paragraphs[0].add_run(h)
		_font(r, 10, True, color=RGBColor(255, 255, 255))
		_shd(hdr[i], C_TH_BG)

	widths = [1.5, 6, 2, 2.5, 2.5, 2]
	for ri, req in enumerate(project.requirements):
		cells = tbl.add_row().cells
		bg = "EBF3FB" if ri % 2 == 0 else "FFFFFF"
		vals = [
			req.req_id or "",
			req.requirement_text or "",
			req.gap_status or "",
			req.mapped_app or "",
			req.mapped_doctype or "",
			str(req.effort_days or 0)
		]
		for ci, val in enumerate(vals):
			cells[ci].text = ""
			r = cells[ci].paragraphs[0].add_run(val)
			_font(r, 9)
			cells[ci].width = Cm(widths[ci])
			if bg != "FFFFFF":
				_shd(cells[ci], bg)

	# Tổng effort
	total_effort = sum(req.effort_days or 0 for req in project.requirements)
	p = doc.add_paragraph()
	r = p.add_run(f"\nTổng effort ước tính: {total_effort} ngày")
	_font(r, 11, True)


# ── Font / style helpers ──────────────────────────────────────────────────────

def _font(run, size=11, bold=False, italic=False, color=None):
	run.font.name = "Times New Roman"
	run.font.size = Pt(size)
	run.font.bold = bold
	run.font.italic = italic
	if color:
		run.font.color.rgb = color


def _shd(cell, hex_color: str):
	tc = cell._tc
	tcPr = tc.get_or_add_tcPr()
	shd = OxmlElement("w:shd")
	shd.set(qn("w:val"), "clear")
	shd.set(qn("w:color"), "auto")
	shd.set(qn("w:fill"), hex_color)
	tcPr.append(shd)
